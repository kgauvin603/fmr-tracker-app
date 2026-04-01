from __future__ import annotations

from pathlib import Path


def extract_text_from_docx(file_path) -> str:
    """
    Extract plain text from a .docx file preserving paragraph structure.
    Tables are extracted row by row with cells separated by tabs.
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required: pip install python-docx")

    doc = Document(str(file_path))
    parts = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            # Paragraph
            text = "".join(run.text for run in element.iterchildren()
                          if run.tag.endswith("}r"))
            # Get full paragraph text via python-docx
            for para in doc.paragraphs:
                if para._element is element:
                    text = para.text
                    break
            if text.strip():
                parts.append(text.strip())

        elif tag == "tbl":
            # Table — find matching table object
            for table in doc.tables:
                if table._element is element:
                    for row in table.rows:
                        row_text = "\t".join(cell.text.strip() for cell in row.cells)
                        if row_text.strip():
                            parts.append(row_text)
                    parts.append("")  # blank line after table
                    break

    return "\n".join(parts)
